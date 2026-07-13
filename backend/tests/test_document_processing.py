import io
import uuid
from pathlib import Path

import pytest
from docx import Document as WordDocument
from fastapi.testclient import TestClient


DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument."
    "wordprocessingml.document"
)


def create_account_and_get_token(
    client: TestClient,
) -> str:
    """Create a unique test account and return its JWT token."""

    unique_value = uuid.uuid4().hex[:8]

    organization_slug = f"processing-test-{unique_value}"
    email = f"admin-{unique_value}@processing-test.example"
    password = "ProcessingTest!2026"

    registration_response = client.post(
        "/api/v1/register",
        json={
            "organization_name": "Processing Test Company",
            "organization_slug": organization_slug,
            "full_name": "Processing Test Administrator",
            "email": email,
            "password": password,
        },
    )

    assert registration_response.status_code == 201

    login_response = client.post(
        "/api/v1/auth/login",
        json={
            "organization_slug": organization_slug,
            "email": email,
            "password": password,
        },
    )

    assert login_response.status_code == 200

    return login_response.json()["access_token"]


def create_test_docx() -> bytes:
    """Create a small valid DOCX file in memory."""

    memory_file = io.BytesIO()

    document = WordDocument()
    document.add_heading("EnterpriseMind AI Handbook", level=1)
    document.add_paragraph(
        "EnterpriseMind AI securely processes private company documents."
    )
    document.add_paragraph(
        "Employees receive twenty days of annual leave each year."
    )
    document.add_paragraph(
        "Only authorised organisation members can access company files."
    )

    document.save(memory_file)

    return memory_file.getvalue()


def test_process_document_creates_chunks(
    client: TestClient,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Processing a DOCX should create stored text chunks."""

    test_upload_directory = tmp_path / "uploads"

    monkeypatch.setattr(
        "backend.app.api.documents.UPLOAD_ROOT",
        test_upload_directory,
    )

    access_token = create_account_and_get_token(client)

    upload_response = client.post(
        "/api/v1/documents/upload",
        headers={
            "Authorization": f"Bearer {access_token}",
        },
        files={
            "file": (
                "employee-handbook.docx",
                create_test_docx(),
                DOCX_CONTENT_TYPE,
            ),
        },
    )

    assert upload_response.status_code == 201

    document_id = upload_response.json()["id"]

    processing_response = client.post(
        f"/api/v1/documents/{document_id}/process",
        headers={
            "Authorization": f"Bearer {access_token}",
        },
    )

    assert processing_response.status_code == 200

    processing_data = processing_response.json()

    assert processing_data["document_id"] == document_id
    assert processing_data["status"] == "ready"
    assert processing_data["chunk_count"] >= 1

    chunks_response = client.get(
        f"/api/v1/documents/{document_id}/chunks",
        headers={
            "Authorization": f"Bearer {access_token}",
        },
    )

    assert chunks_response.status_code == 200

    chunks = chunks_response.json()

    assert len(chunks) == processing_data["chunk_count"]
    assert chunks[0]["chunk_index"] == 0
    assert chunks[0]["page_number"] is None
    assert chunks[0]["character_count"] > 0
    assert "EnterpriseMind AI" in chunks[0]["content"]


def test_chunks_endpoint_requires_authentication(
    client: TestClient,
) -> None:
    """Users without a token must not access chunks."""

    document_id = uuid.uuid4()

    response = client.get(
        f"/api/v1/documents/{document_id}/chunks"
    )

    assert response.status_code == 401


def test_chunks_endpoint_protects_organization_data(
    client: TestClient,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """One organisation must not access another organisation's chunks."""

    test_upload_directory = tmp_path / "uploads"

    monkeypatch.setattr(
        "backend.app.api.documents.UPLOAD_ROOT",
        test_upload_directory,
    )

    first_token = create_account_and_get_token(client)
    second_token = create_account_and_get_token(client)

    upload_response = client.post(
        "/api/v1/documents/upload",
        headers={
            "Authorization": f"Bearer {first_token}",
        },
        files={
            "file": (
                "private-handbook.docx",
                create_test_docx(),
                DOCX_CONTENT_TYPE,
            ),
        },
    )

    assert upload_response.status_code == 201

    document_id = upload_response.json()["id"]

    processing_response = client.post(
        f"/api/v1/documents/{document_id}/process",
        headers={
            "Authorization": f"Bearer {first_token}",
        },
    )

    assert processing_response.status_code == 200

    second_organization_response = client.get(
        f"/api/v1/documents/{document_id}/chunks",
        headers={
            "Authorization": f"Bearer {second_token}",
        },
    )

    assert second_organization_response.status_code == 404
    assert second_organization_response.json() == {
        "detail": "Document not found."
    }