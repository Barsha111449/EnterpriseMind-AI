import io
import uuid
from pathlib import Path

import pytest
from docx import Document as WordDocument
from fastapi.testclient import TestClient


DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument."
    "wordprocessingml.document"
)

TEST_EMBEDDING = [1.0] + [0.0] * 383


def create_account_and_get_token(
    client: TestClient,
) -> str:
    """Create a unique organisation account and return its token."""

    unique_value = uuid.uuid4().hex[:8]

    organization_slug = f"search-test-{unique_value}"
    email = f"admin-{unique_value}@search-test.example"
    password = "SearchTest!2026"

    registration_response = client.post(
        "/api/v1/register",
        json={
            "organization_name": "Search Test Company",
            "organization_slug": organization_slug,
            "full_name": "Search Test Administrator",
            "email": email,
            "password": password,
        },
    )

    assert registration_response.status_code == 201

    login_response = client.post(
        "/api/v1/auth/login",
        json={
            "organization_slug": organization_slug,
            "email": email,
            "password": password,
        },
    )

    assert login_response.status_code == 200

    return login_response.json()["access_token"]


def create_test_docx(
    title: str,
    content: str,
) -> bytes:
    """Create a valid DOCX document in memory."""

    memory_file = io.BytesIO()

    document = WordDocument()
    document.add_heading(title, level=1)
    document.add_paragraph(content)
    document.add_paragraph(
        "EnterpriseMind AI protects private organisation data."
    )

    document.save(memory_file)

    return memory_file.getvalue()


def configure_fake_embeddings(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Replace the real embedding model with deterministic vectors."""

    monkeypatch.setattr(
        "backend.app.api.documents.generate_embeddings",
        lambda texts: [
            TEST_EMBEDDING.copy()
            for _ in texts
        ],
    )

    monkeypatch.setattr(
        "backend.app.api.search.generate_embeddings",
        lambda texts: [
            TEST_EMBEDDING.copy()
            for _ in texts
        ],
    )


def upload_and_process_document(
    client: TestClient,
    access_token: str,
    filename: str,
    content: bytes,
) -> str:
    """Upload and process one document."""

    headers = {
        "Authorization": f"Bearer {access_token}",
    }

    upload_response = client.post(
        "/api/v1/documents/upload",
        headers=headers,
        files={
            "file": (
                filename,
                content,
                DOCX_CONTENT_TYPE,
            ),
        },
    )

    assert upload_response.status_code == 201

    document_id = upload_response.json()["id"]

    processing_response = client.post(
        f"/api/v1/documents/{document_id}/process",
        headers=headers,
    )

    assert processing_response.status_code == 200
    assert processing_response.json()["status"] == "ready"

    return document_id


def test_semantic_search_requires_authentication(
    client: TestClient,
) -> None:
    """Users without a token must not use semantic search."""

    response = client.post(
        "/api/v1/search/semantic",
        json={
            "query": "What is the leave policy?",
            "top_k": 5,
        },
    )

    assert response.status_code == 401


def test_semantic_search_returns_document_chunks(
    client: TestClient,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Semantic search should return relevant stored chunks."""

    test_upload_directory = tmp_path / "uploads"

    monkeypatch.setattr(
        "backend.app.api.documents.UPLOAD_ROOT",
        test_upload_directory,
    )

    configure_fake_embeddings(monkeypatch)

    access_token = create_account_and_get_token(client)

    upload_and_process_document(
        client=client,
        access_token=access_token,
        filename="employee-handbook.docx",
        content=create_test_docx(
            title="Employee Handbook",
            content=(
                "Employees receive twenty days of annual leave "
                "during each calendar year."
            ),
        ),
    )

    response = client.post(
        "/api/v1/search/semantic",
        headers={
            "Authorization": f"Bearer {access_token}",
        },
        json={
            "query": "How many annual leave days are provided?",
            "top_k": 5,
        },
    )

    assert response.status_code == 200

    response_data = response.json()

    assert response_data["query"] == (
        "How many annual leave days are provided?"
    )

    assert response_data["result_count"] >= 1
    assert len(response_data["results"]) >= 1

    first_result = response_data["results"][0]

    assert first_result["original_filename"] == (
        "employee-handbook.docx"
    )

    assert "content" in first_result
    assert "similarity_score" in first_result
    assert first_result["similarity_score"] == pytest.approx(
        1.0
    )


def test_semantic_search_respects_top_k(
    client: TestClient,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """The top_k value should limit returned search results."""

    test_upload_directory = tmp_path / "uploads"

    monkeypatch.setattr(
        "backend.app.api.documents.UPLOAD_ROOT",
        test_upload_directory,
    )

    configure_fake_embeddings(monkeypatch)

    access_token = create_account_and_get_token(client)

    upload_and_process_document(
        client=client,
        access_token=access_token,
        filename="leave-policy.docx",
        content=create_test_docx(
            title="Leave Policy",
            content="Employees receive annual leave.",
        ),
    )

    upload_and_process_document(
        client=client,
        access_token=access_token,
        filename="security-policy.docx",
        content=create_test_docx(
            title="Security Policy",
            content="Company files require authorised access.",
        ),
    )

    response = client.post(
        "/api/v1/search/semantic",
        headers={
            "Authorization": f"Bearer {access_token}",
        },
        json={
            "query": "What company policies are available?",
            "top_k": 1,
        },
    )

    assert response.status_code == 200

    response_data = response.json()

    assert response_data["result_count"] == 1
    assert len(response_data["results"]) == 1


def test_semantic_search_protects_organization_data(
    client: TestClient,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """One organisation must not search another organisation's data."""

    test_upload_directory = tmp_path / "uploads"

    monkeypatch.setattr(
        "backend.app.api.documents.UPLOAD_ROOT",
        test_upload_directory,
    )

    configure_fake_embeddings(monkeypatch)

    first_token = create_account_and_get_token(client)
    second_token = create_account_and_get_token(client)

    upload_and_process_document(
        client=client,
        access_token=first_token,
        filename="first-company-private.docx",
        content=create_test_docx(
            title="First Company Private Policy",
            content=(
                "The first company has a confidential "
                "internal leave policy."
            ),
        ),
    )

    upload_and_process_document(
        client=client,
        access_token=second_token,
        filename="second-company-private.docx",
        content=create_test_docx(
            title="Second Company Private Policy",
            content=(
                "The second company has confidential "
                "financial information."
            ),
        ),
    )

    response = client.post(
        "/api/v1/search/semantic",
        headers={
            "Authorization": f"Bearer {first_token}",
        },
        json={
            "query": "Show confidential company information.",
            "top_k": 20,
        },
    )

    assert response.status_code == 200

    results = response.json()["results"]

    assert len(results) >= 1

    returned_filenames = {
        result["original_filename"]
        for result in results
    }

    assert "first-company-private.docx" in returned_filenames

    assert (
        "second-company-private.docx"
        not in returned_filenames
    )
def test_keyword_search_returns_matching_chunks(
    client: TestClient,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Keyword search should return chunks containing exact words."""

    test_upload_directory = tmp_path / "uploads"

    monkeypatch.setattr(
        "backend.app.api.documents.UPLOAD_ROOT",
        test_upload_directory,
    )

    configure_fake_embeddings(monkeypatch)

    access_token = create_account_and_get_token(client)

    upload_and_process_document(
        client=client,
        access_token=access_token,
        filename="technical-skills.docx",
        content=create_test_docx(
            title="Technical Skills",
            content=(
                "The employee has Python and machine learning "
                "development experience."
            ),
        ),
    )

    response = client.post(
        "/api/v1/search/keyword",
        headers={
            "Authorization": f"Bearer {access_token}",
        },
        json={
            "query": "Python machine learning",
            "top_k": 5,
        },
    )

    assert response.status_code == 200

    response_data = response.json()

    assert response_data["query"] == (
        "Python machine learning"
    )

    assert response_data["result_count"] >= 1

    matching_result = next(
        result
        for result in response_data["results"]
        if result["original_filename"]
        == "technical-skills.docx"
    )

    assert matching_result["keyword_score"] > 0
    assert "Python" in matching_result["content"]


def test_keyword_search_requires_authentication(
    client: TestClient,
) -> None:
    """Unauthenticated users must not use keyword search."""

    response = client.post(
        "/api/v1/search/keyword",
        json={
            "query": "Python",
            "top_k": 5,
        },
    )

    assert response.status_code == 401


def test_hybrid_search_requires_authentication(
    client: TestClient,
) -> None:
    """Unauthenticated users must not use hybrid search."""

    response = client.post(
        "/api/v1/search/hybrid",
        json={
            "query": "company policy",
            "top_k": 5,
        },
    )

    assert response.status_code == 401


def test_hybrid_search_combines_search_results(
    client: TestClient,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Hybrid search should combine semantic and keyword results."""

    test_upload_directory = tmp_path / "uploads"

    monkeypatch.setattr(
        "backend.app.api.documents.UPLOAD_ROOT",
        test_upload_directory,
    )

    configure_fake_embeddings(monkeypatch)

    access_token = create_account_and_get_token(client)

    upload_and_process_document(
        client=client,
        access_token=access_token,
        filename="ai-skills.docx",
        content=create_test_docx(
            title="AI Engineering Skills",
            content=(
                "The engineer uses Python for machine learning "
                "and artificial intelligence projects."
            ),
        ),
    )

    response = client.post(
        "/api/v1/search/hybrid",
        headers={
            "Authorization": f"Bearer {access_token}",
        },
        json={
            "query": "Python machine learning",
            "top_k": 5,
        },
    )

    assert response.status_code == 200

    response_data = response.json()

    assert response_data["query"] == (
        "Python machine learning"
    )

    assert response_data["result_count"] >= 1

    matching_result = next(
        result
        for result in response_data["results"]
        if result["original_filename"]
        == "ai-skills.docx"
    )

    assert matching_result["semantic_score"] is not None
    assert matching_result["keyword_score"] is not None
    assert matching_result["hybrid_score"] > 0


def test_hybrid_search_respects_top_k(
    client: TestClient,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Hybrid search must respect the requested result limit."""

    test_upload_directory = tmp_path / "uploads"

    monkeypatch.setattr(
        "backend.app.api.documents.UPLOAD_ROOT",
        test_upload_directory,
    )

    configure_fake_embeddings(monkeypatch)

    access_token = create_account_and_get_token(client)

    upload_and_process_document(
        client=client,
        access_token=access_token,
        filename="python-policy.docx",
        content=create_test_docx(
            title="Python Policy",
            content="Python is used for company applications.",
        ),
    )

    upload_and_process_document(
        client=client,
        access_token=access_token,
        filename="machine-learning-policy.docx",
        content=create_test_docx(
            title="Machine Learning Policy",
            content=(
                "Machine learning is used for company systems."
            ),
        ),
    )

    response = client.post(
        "/api/v1/search/hybrid",
        headers={
            "Authorization": f"Bearer {access_token}",
        },
        json={
            "query": "company systems",
            "top_k": 1,
        },
    )

    assert response.status_code == 200

    response_data = response.json()

    assert response_data["result_count"] == 1
    assert len(response_data["results"]) == 1


def test_hybrid_search_protects_organization_data(
    client: TestClient,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Hybrid search must not return another organisation's data."""

    test_upload_directory = tmp_path / "uploads"

    monkeypatch.setattr(
        "backend.app.api.documents.UPLOAD_ROOT",
        test_upload_directory,
    )

    configure_fake_embeddings(monkeypatch)

    first_token = create_account_and_get_token(client)
    second_token = create_account_and_get_token(client)

    upload_and_process_document(
        client=client,
        access_token=first_token,
        filename="first-company-policy.docx",
        content=create_test_docx(
            title="First Company Policy",
            content=(
                "The confidential Phoenix policy belongs only "
                "to the first company."
            ),
        ),
    )

    upload_and_process_document(
        client=client,
        access_token=second_token,
        filename="second-company-policy.docx",
        content=create_test_docx(
            title="Second Company Policy",
            content=(
                "The second company stores confidential "
                "financial information."
            ),
        ),
    )

    response = client.post(
        "/api/v1/search/hybrid",
        headers={
            "Authorization": f"Bearer {first_token}",
        },
        json={
            "query": "confidential Phoenix policy",
            "top_k": 20,
        },
    )

    assert response.status_code == 200

    returned_filenames = {
        result["original_filename"]
        for result in response.json()["results"]
    }

    assert "first-company-policy.docx" in returned_filenames

    assert (
        "second-company-policy.docx"
        not in returned_filenames
    )