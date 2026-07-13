import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document as WordDocument
from pypdf import PdfReader


MAX_CHUNK_CHARACTERS = 1200
CHUNK_OVERLAP_CHARACTERS = 200


@dataclass(frozen=True, slots=True)
class ExtractedChunk:
    """A piece of extracted document text."""

    content: str
    page_number: int | None


def normalize_text(text: str) -> str:
    """Remove unnecessary spaces and blank lines."""

    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


def split_text(
    text: str,
    max_characters: int = MAX_CHUNK_CHARACTERS,
    overlap_characters: int = CHUNK_OVERLAP_CHARACTERS,
) -> list[str]:
    """Split text into overlapping chunks."""

    if max_characters <= 0:
        raise ValueError(
            "Maximum chunk size must be greater than zero."
        )

    if overlap_characters < 0:
        raise ValueError(
            "Chunk overlap cannot be negative."
        )

    if overlap_characters >= max_characters:
        raise ValueError(
            "Chunk overlap must be smaller than chunk size."
        )

    cleaned_text = normalize_text(text)

    if not cleaned_text:
        return []

    chunks: list[str] = []
    start = 0

    while start < len(cleaned_text):
        end = min(
            start + max_characters,
            len(cleaned_text),
        )

        if end < len(cleaned_text):
            candidate = cleaned_text[start:end]

            boundary = max(
                candidate.rfind("\n\n"),
                candidate.rfind(". "),
                candidate.rfind(" "),
            )

            if boundary >= max_characters // 2:
                end = start + boundary + 1

        chunk = cleaned_text[start:end].strip()

        if chunk:
            chunks.append(chunk)

        if end >= len(cleaned_text):
            break

        start = max(
            end - overlap_characters,
            start + 1,
        )

    return chunks


def extract_pdf_chunks(
    file_path: Path,
) -> list[ExtractedChunk]:
    """Extract page-aware text chunks from a PDF."""

    reader = PdfReader(str(file_path))
    extracted_chunks: list[ExtractedChunk] = []

    for page_number, page in enumerate(
        reader.pages,
        start=1,
    ):
        page_text = page.extract_text() or ""

        for chunk_text in split_text(page_text):
            extracted_chunks.append(
                ExtractedChunk(
                    content=chunk_text,
                    page_number=page_number,
                )
            )

    return extracted_chunks


def extract_docx_chunks(
    file_path: Path,
) -> list[ExtractedChunk]:
    """Extract text chunks from a DOCX document."""

    document = WordDocument(str(file_path))
    text_parts: list[str] = []

    for paragraph in document.paragraphs:
        paragraph_text = normalize_text(paragraph.text)

        if paragraph_text:
            text_parts.append(paragraph_text)

    for table in document.tables:
        for row in table.rows:
            cell_values = [
                normalize_text(cell.text)
                for cell in row.cells
            ]

            non_empty_values = [
                value
                for value in cell_values
                if value
            ]

            if non_empty_values:
                text_parts.append(
                    " | ".join(non_empty_values)
                )

    full_text = "\n\n".join(text_parts)

    return [
        ExtractedChunk(
            content=chunk_text,
            page_number=None,
        )
        for chunk_text in split_text(full_text)
    ]


def extract_document_chunks(
    file_path: Path,
) -> list[ExtractedChunk]:
    """Extract chunks according to the document type."""

    extension = file_path.suffix.lower()

    if extension == ".pdf":
        return extract_pdf_chunks(file_path)

    if extension == ".docx":
        return extract_docx_chunks(file_path)

    raise ValueError(
        "Only PDF and DOCX documents can be processed."
    )